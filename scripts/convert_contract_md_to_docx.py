from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
CONTRACT_DIR = ROOT / "docs" / "contract"
DOCS = [
    CONTRACT_DIR / "ICP备案材料及流程.md",
    CONTRACT_DIR / "研发范围及功能描述.md",
    CONTRACT_DIR / "阿里云资源清单及价格.md",
]


def set_run_font(run, name="Microsoft YaHei", size=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_style_font(style, name="Microsoft YaHei", size=None, bold=None):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(clean_inline(text))
    set_run_font(run, size=9.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def clean_inline(text):
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("\\|", "|")
    return text.strip()


def split_table_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [clean_inline(part.strip()) for part in line.split("|")]


def is_table_separator(line):
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def add_table(document, rows):
    if len(rows) < 2:
        return
    header = split_table_row(rows[0])
    body = [split_table_row(row) for row in rows[2:] if row.strip()]
    column_count = len(header)
    table = document.add_table(rows=1, cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for idx, text in enumerate(header):
        set_cell_text(table.rows[0].cells[idx], text, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "D9EAF7")

    for row in body:
        cells = table.add_row().cells
        for idx in range(column_count):
            set_cell_text(cells[idx], row[idx] if idx < len(row) else "")

    document.add_paragraph()


def add_paragraph(document, text, style=None, bullet=False, quote=False):
    paragraph = document.add_paragraph(style=style)
    if bullet:
        paragraph.style = document.styles["List Bullet"]
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    if quote:
        paragraph.paragraph_format.left_indent = Cm(0.6)
    run = paragraph.add_run(clean_inline(text))
    set_run_font(run, size=10.5)
    if quote:
        run.font.color.rgb = RGBColor(89, 89, 89)
    return paragraph


def configure_document(document):
    section = document.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    set_style_font(document.styles["Normal"], size=10.5)
    set_style_font(document.styles["Title"], size=18, bold=True)
    set_style_font(document.styles["Heading 1"], size=15, bold=True)
    set_style_font(document.styles["Heading 2"], size=13, bold=True)
    set_style_font(document.styles["Heading 3"], size=11.5, bold=True)
    set_style_font(document.styles["List Bullet"], size=10.5)


def convert(path):
    document = Document()
    configure_document(document)
    lines = path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith("```"):
            if in_code:
                add_paragraph(document, "\n".join(code_lines), style=None)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.strip().startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(document, table_lines)
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            text = clean_inline(heading.group(2))
            if level == 1:
                paragraph = document.add_paragraph(style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(text)
                set_run_font(run, size=18, bold=True)
            else:
                style = f"Heading {min(level - 1, 3)}"
                paragraph = document.add_paragraph(style=style)
                run = paragraph.add_run(text)
                set_run_font(run, size={2: 15, 3: 13}.get(level, 11.5), bold=True)
            i += 1
            continue

        if line.lstrip().startswith(">"):
            add_paragraph(document, line.lstrip()[1:].strip(), quote=True)
            i += 1
            continue

        bullet = re.match(r"^\s*[-*]\s+(.+)$", line)
        if bullet:
            add_paragraph(document, bullet.group(1), bullet=True)
            i += 1
            continue

        numbered = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if numbered:
            add_paragraph(document, numbered.group(1))
            i += 1
            continue

        add_paragraph(document, line)
        i += 1

    output = path.with_suffix(".docx")
    document.save(output)
    return output


def main():
    for path in DOCS:
        output = convert(path)
        print(output)


if __name__ == "__main__":
    main()
